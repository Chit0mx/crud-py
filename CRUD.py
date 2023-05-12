import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk


import sqlite3

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch

from docx import Document


# Crear la conexión a la base de datos SQLite
conexion = sqlite3.connect("participantes.db")
cursor = conexion.cursor()

# Crear la tabla de participantes si no existe
cursor.execute("""
    CREATE TABLE IF NOT EXISTS participantes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT,
        apellido_paterno TEXT,
        apellido_materno TEXT,
        edad INTEGER,
        sexo TEXT,
        escuela TEXT,
        direccion TEXT,
        curp TEXT,
        categoria TEXT
    )
""")
conexion.commit()

# Función para agregar un nuevo participante a la base de datos
def agregar_participante():
    nombre = entry_nombre.get()
    apellido_paterno = entry_apellido_paterno.get()
    apellido_materno = entry_apellido_materno.get()
    edad = entry_edad.get()
    sexo = entry_sexo.get()
    escuela = entry_escuela.get()
    direccion = entry_direccion.get()
    curp = entry_curp.get()
    categoria = entry_categoria.get()

    if nombre and apellido_paterno and edad and curp and categoria:
        cursor.execute("""
            INSERT INTO participantes (nombre, apellido_paterno, apellido_materno, edad, sexo, escuela, direccion, curp, categoria)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (nombre, apellido_paterno, apellido_materno, edad, sexo, escuela, direccion, curp, categoria))
        conexion.commit()
        messagebox.showinfo("Éxito", "El participante ha sido agregado correctamente.")
        limpiar_campos()
        actualizar_lista()
    else:
        messagebox.showerror("Error", "Por favor, completa los campos obligatorios.")

# Función para actualizar un participante existente
def actualizar_participante():
    # Obtener el índice seleccionado de la lista
    indice = lista.curselection()
    if not indice:
        messagebox.showwarning("Advertencia", "Selecciona un participante de la lista.")
        return

    # Obtener los datos del participante seleccionado
    participante_seleccionado = lista.get(indice)
    id_participante = participante_seleccionado[0]

    # Obtener los valores de los campos de entrada
    nombre = entry_nombre.get()
    apellido_paterno = entry_apellido_paterno.get()
    apellido_materno = entry_apellido_materno.get()
    edad = entry_edad.get()
    sexo = entry_sexo.get()
    escuela = entry_escuela.get()
    direccion = entry_direccion.get()
    curp = entry_curp.get()
    categoria = entry_categoria.get()

    # Validar que los campos no estén vacíos
    if not nombre or not apellido_paterno or not apellido_materno or not edad or not sexo or not escuela or not direccion or not curp or not categoria:
        messagebox.showwarning("Advertencia", "Completa todos los campos.")
        return

    # Actualizar los datos del participante en la base de datos
    cursor.execute("UPDATE participantes SET nombre=?, apellido_paterno=?, apellido_materno=?, edad=?, sexo=?, escuela=?, direccion=?, curp=?, categoria=? WHERE id=?", (nombre, apellido_paterno, apellido_materno, edad, sexo, escuela, direccion, curp, categoria, id_participante))
    conexion.commit()

    # Limpiar los campos de entrada
    entry_nombre.delete(0, tk.END)
    entry_apellido_paterno.delete(0, tk.END)
    entry_apellido_materno.delete(0, tk.END)
    entry_edad.delete(0, tk.END)
    entry_sexo.delete(0, tk.END)
    entry_escuela.delete(0, tk.END)
    entry_direccion.delete(0, tk.END)
    entry_curp.delete(0, tk.END)
    entry_categoria.delete(0, tk.END)

    # Actualizar la lista de participantes
    actualizar_lista()

    messagebox.showinfo("Éxito", "Participante actualizado correctamente.")

def eliminar_participante():
    indice = lista.curselection()
    if indice:
        id_participante = lista.get(indice)[0]
        respuesta = messagebox.askquestion("Confirmar", "¿Estás seguro de eliminar este participante?")
        if respuesta == "yes":
            cursor.execute("DELETE FROM participantes WHERE id=?", (id_participante,))
            conexion.commit()
            messagebox.showinfo("Éxito", "El participante ha sido eliminado correctamente.")
            limpiar_campos()
            actualizar_lista()

def mostrar_datos(event):
    # Obtener el índice del registro seleccionado
    indice = lista.curselection()

    if indice:
        # Obtener los datos del registro seleccionado
        id_participante = lista.get(indice)[0]

        # Obtener los datos completos del registro desde la base de datos
        cursor.execute("SELECT * FROM participantes WHERE id=?", (id_participante,))
        datos_registro = cursor.fetchone()

        # Mostrar los datos en los campos del formulario
        limpiar_campos()
        entry_nombre.insert(tk.END, datos_registro[1])
        entry_apellido_paterno.insert(tk.END, datos_registro[2])
        entry_apellido_materno.insert(tk.END, datos_registro[3])
        entry_edad.insert(tk.END, datos_registro[4])
        entry_sexo.insert(tk.END, datos_registro[5])
        entry_escuela.insert(tk.END, datos_registro[6])
        entry_direccion.insert(tk.END, datos_registro[7])
        entry_curp.insert(tk.END, datos_registro[8])
        entry_categoria.insert(tk.END, datos_registro[9])

def buscar_registro():
    # Obtener el nombre buscado
    nombre_busqueda = entry_buscar.get().lower()

    # Limpiar la lista
    lista.delete(0, tk.END)

    # Realizar la consulta a la base de datos
    cursor.execute("SELECT * FROM participantes WHERE nombre LIKE ?", ('%' + nombre_busqueda + '%',))
    registros_encontrados = cursor.fetchall()

    # Agregar registros encontrados a la lista
    for registro in registros_encontrados:
        lista.insert(tk.END, registro)

def generar_pdf():
    indice = lista.curselection()
    if indice:
        participante = lista.get(indice)
        nombre_pdf = f"{participante[0]}_{participante[1]}_{participante[2]}.pdf"
        c = canvas.Canvas(nombre_pdf)
        c.setFont("Helvetica", 12)
        c.drawString(50, 750, "Datos del participante:")
        c.drawString(50, 700, f"Nombre: {participante[1]} {participante[2]} {participante[3]}")
        c.drawString(50, 680, f"Edad: {participante[4]}")
        c.drawString(50, 660, f"Sexo: {participante[5]}")
        c.drawString(50, 640, f"Escuela: {participante[6]}")
        c.drawString(50, 620, f"Dirección: {participante[7]}")
        c.drawString(50, 600, f"CURP: {participante[8]}")
        c.drawString(50, 580, f"Categoría: {participante[9]}")
        c.save()
        messagebox.showinfo("Éxito", "El PDF ha sido generado correctamente.")

def exportar_registro():
    # Obtener el índice seleccionado de la lista
    indice = lista.curselection()
    if not indice:
        messagebox.showwarning("Advertencia", "Selecciona un participante de la lista.")
        return

    # Obtener los datos del participante seleccionado
    participante_seleccionado = lista.get(indice)

    # Crear el contenido del archivo de texto
    contenido = f"Nombre: {participante_seleccionado[1]}\n"
    contenido += f"Apellido Paterno: {participante_seleccionado[2]}\n"
    contenido += f"Apellido Materno: {participante_seleccionado[3]}\n"
    contenido += f"Edad: {participante_seleccionado[4]}\n"
    contenido += f"Sexo: {participante_seleccionado[5]}\n"
    contenido += f"Escuela: {participante_seleccionado[6]}\n"
    contenido += f"Dirección: {participante_seleccionado[7]}\n"
    contenido += f"CURP: {participante_seleccionado[8]}\n"
    contenido += f"Categoría: {participante_seleccionado[9]}"

    # Guardar el archivo de texto
    archivo = filedialog.asksaveasfile(defaultextension=".txt", filetypes=[("Archivos de texto", "*.txt")])
    if archivo is None:
        return

    archivo.write(contenido)
    archivo.close()

    messagebox.showinfo("Éxito", "Registro exportado correctamente.")


def exportar_registro_docx():
    # Obtener el índice seleccionado de la lista
    indice = lista.curselection()
    if not indice:
        messagebox.showwarning("Advertencia", "Selecciona un participante de la lista.")
        return

    # Obtener los datos del participante seleccionado
    participante_seleccionado = lista.get(indice)

    # Crear un nuevo documento de Word
    doc = Document()

    # Agregar los datos del participante al documento
    doc.add_paragraph(f"Nombre: {participante_seleccionado[1]}")
    doc.add_paragraph(f"Apellido Paterno: {participante_seleccionado[2]}")
    doc.add_paragraph(f"Apellido Materno: {participante_seleccionado[3]}")
    doc.add_paragraph(f"Edad: {participante_seleccionado[4]}")
    doc.add_paragraph(f"Sexo: {participante_seleccionado[5]}")
    doc.add_paragraph(f"Escuela: {participante_seleccionado[6]}")
    doc.add_paragraph(f"Dirección: {participante_seleccionado[7]}")
    doc.add_paragraph(f"CURP: {participante_seleccionado[8]}")
    doc.add_paragraph(f"Categoría: {participante_seleccionado[9]}")

    # Guardar el archivo DOCX
    archivo = filedialog.asksaveasfile(defaultextension=".docx", filetypes=[("Archivos de Word", "*.docx")])
    if archivo is None:
        return

    doc.save(archivo.name)
    archivo.close()

    messagebox.showinfo("Éxito", "Registro exportado correctamente como archivo DOCX.")

def generar_reconocimiento():
    # Obtener el índice seleccionado de la lista
    indice = lista.curselection()
    if not indice:
        messagebox.showwarning("Advertencia", "Selecciona un participante de la lista.")
        return

    # Obtener los datos del participante seleccionado
    participante_seleccionado = lista.get(indice)

    # Crear el nombre del archivo PDF
    nombre_archivo = f"Reconocimiento_{participante_seleccionado[1]}_{participante_seleccionado[2]}.pdf"

    # Crear el documento PDF
    c = canvas.Canvas(nombre_archivo, pagesize=letter)

    # Definir el estilo de la página
    c.setFillColor(colors.lightblue)
    c.rect(0, 0, letter[0], letter[1], fill=True)

    # Agregar el título
    c.setFont("Helvetica-Bold", 20)
    c.setFillColor(colors.white)
    c.drawString(1.5 * inch, 9 * inch, "RECONOCIMIENTO")

    # Agregar el nombre del participante
    c.setFont("Helvetica-Bold", 16)
    c.setFillColor(colors.white)
    c.drawString(1.5 * inch, 7 * inch, f"Se otorga este reconocimiento a:")
    c.setFont("Helvetica-Bold", 20)
    c.setFillColor(colors.black)
    c.drawString(1.5 * inch, 6 * inch, f"{participante_seleccionado[1]} {participante_seleccionado[2]}")

    # Agregar el mensaje del reconocimiento
    mensaje = "Por su destacada participación en el evento."
    c.setFont("Helvetica", 14)
    c.setFillColor(colors.black)
    c.drawString(1.5 * inch, 5 * inch, mensaje)

    # Agregar la firma del organizador
    firma = "Organización del Evento"
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(colors.black)
    c.drawString(1.5 * inch, 3 * inch, firma)

    # Guardar y cerrar el documento PDF
    c.showPage()
    c.save()

    messagebox.showinfo("Éxito", f"Reconocimiento generado correctamente. Se guardó como {nombre_archivo}.")

def limpiar_campos():
    entry_nombre.delete(0, tk.END)
    entry_apellido_paterno.delete(0, tk.END)
    entry_apellido_materno.delete(0, tk.END)
    entry_edad.delete(0, tk.END)
    entry_sexo.delete(0, tk.END)
    entry_escuela.delete(0, tk.END)
    entry_direccion.delete(0, tk.END)
    entry_curp.delete(0, tk.END)
    entry_categoria.delete(0, tk.END)

def actualizar_lista():
    lista.delete(0, tk.END)
    cursor.execute("SELECT * FROM participantes")
    participantes = cursor.fetchall()
    for participante in participantes:
        lista.insert(tk.END, participante)

# Crear la ventana principal
ventana = tk.Tk()
ventana.title("Concurso")

# Crear los campos de entrada
entry_nombre = tk.Entry(ventana)
entry_nombre.grid(row=0, column=1, padx=5, pady=5)

entry_apellido_paterno = tk.Entry(ventana)
entry_apellido_paterno.grid(row=1, column=1, padx=5, pady=5)

entry_apellido_materno = tk.Entry(ventana)
entry_apellido_materno.grid(row=2, column=1, padx=5, pady=5)

entry_edad = tk.Entry(ventana)
entry_edad.grid(row=3, column=1, padx=5, pady=5)

entry_sexo = tk.Entry(ventana)
entry_sexo.grid(row=4, column=1, padx=5, pady=5)

entry_escuela = tk.Entry(ventana)
entry_escuela.grid(row=5, column=1, padx=5, pady=5)

entry_direccion = tk.Entry(ventana)
entry_direccion.grid(row=6, column=1, padx=5, pady=5)

entry_curp = tk.Entry(ventana)
entry_curp.grid(row=7, column=1, padx=5, pady=5)

# entry_categoria = tk.Entry(ventana)
# entry_categoria.grid(row=8, column=1, padx=5, pady=5)

label_categoria = ttk.Label(ventana, text="Categoría:")
label_categoria.grid(row=8, column=0, padx=5, pady=5)
categoria_var = tk.StringVar()
entry_categoria = ttk.Combobox(ventana, textvariable=categoria_var)
entry_categoria["values"] = ("Novatos", "Avanzados", "Libre")
entry_categoria.grid(row=8, column=1, padx=5, pady=5)

# Crear etiquetas para los campos de entrada
label_nombre = tk.Label(ventana, text="Nombre:")
label_nombre.grid(row=0, column=0, padx=5, pady=5, sticky="e")

label_apellido_paterno = tk.Label(ventana, text="Apellido Paterno:")
label_apellido_paterno.grid(row=1, column=0, padx=5, pady=5, sticky="e")

label_apellido_materno = tk.Label(ventana, text="Apellido Materno:")
label_apellido_materno.grid(row=2, column=0, padx=5, pady=5, sticky="e")

label_edad = tk.Label(ventana, text="Edad:")
label_edad.grid(row=3, column=0, padx=5, pady=5, sticky="e")

label_sexo = tk.Label(ventana, text="Sexo:")
label_sexo.grid(row=4, column=0, padx=5, pady=5, sticky="e")

label_escuela = tk.Label(ventana, text="Escuela:")
label_escuela.grid(row=5, column=0, padx=5, pady=5, sticky="e")

label_direccion = tk.Label(ventana, text="Dirección:")
label_direccion.grid(row=6, column=0, padx=5, pady=5, sticky="e")

label_curp = tk.Label(ventana, text="CURP:")
label_curp.grid(row=7, column=0, padx=5, pady=5, sticky="e")

label_categoria = tk.Label(ventana, text="Categoría:")
label_categoria.grid(row=8, column=0, padx=5, pady=5, sticky="e")

# Crear la lista de participantes
lista = tk.Listbox(ventana, width=50)
lista.grid(row=0, column=2, rowspan=9, padx=5, pady=5)

# Asociar la función mostrar_datos al evento de selección en la lista
lista.bind("<<ListboxSelect>>", mostrar_datos)

# Crear los botones
frame_botones = tk.Frame(ventana)
frame_botones.grid(row=9, column=0, columnspan=4, padx=5, pady=5)

boton_agregar = tk.Button(frame_botones, text="Agregar", command=agregar_participante)
boton_agregar.grid(row=0, column=0, padx=5)

boton_actualizar = tk.Button(frame_botones, text="Actualizar", command=actualizar_participante)
boton_actualizar.grid(row=0, column=1, padx=5)

boton_eliminar = tk.Button(frame_botones, text="Eliminar", command=eliminar_participante)
boton_eliminar.grid(row=0, column=2, padx=5)

boton_eliminar = tk.Button(frame_botones, text="Limpiar", command=limpiar_campos)
boton_eliminar.grid(row=0, column=3, padx=5)

boton_pdf = tk.Button(frame_botones, text="Exportar registro .pdf", command=generar_pdf)
boton_pdf.grid(row=0, column=4, padx=5)

boton_exportar = tk.Button(ventana, text="Exportar registro .txt", command=exportar_registro)
boton_exportar.grid(row=9, column=0,padx=5, pady=5)

boton_exportar_docx = tk.Button(ventana, text="Exportar registro .docx", command=exportar_registro_docx)
boton_exportar_docx.grid(row=10, column=0, columnspan=3, padx=5, pady=5)

boton_generar_reconocimiento = tk.Button(ventana, text="Generar Reconocimiento", command=generar_reconocimiento)
boton_generar_reconocimiento.grid(row=11, column=0, columnspan=3, padx=5, pady=5)

# Crear la etiqueta y el campo de búsqueda
label_buscar = ttk.Label(ventana, text="Buscar por nombre:")
label_buscar.grid(row=0, column=3, padx=5, pady=5)
entry_buscar = ttk.Entry(ventana)
entry_buscar.grid(row=0, column=4, padx=5, pady=5)
boton_buscar = ttk.Button(ventana, text="Buscar", command=buscar_registro)
boton_buscar.grid(row=0, column=5, padx=5, pady=5)

# Actualizar la lista inicial de participantes
actualizar_lista()

# Ejecutar la interfaz gráfica
ventana.mainloop()

# Cerrar la conexión a la base de datos al finalizar
conexion.close()
